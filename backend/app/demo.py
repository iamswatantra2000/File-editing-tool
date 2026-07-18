"""
Offline demo — runs the full pie-chart flow in mock mode without a server.

Usage:
    cd backend
    python -m app.demo

Produces: demo_output/edited_report.docx
"""

import json
import os
import shutil
import sys
from pathlib import Path

# Force mock mode so no API key is required
os.environ.setdefault("AGENT_MODE", "mock")

from docx import Document as DocxDocument

from app.docops.outline import build_outline
from app.docops.session import DocumentSession
from app.agent.orchestrator import run_agent


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

DEMO_DIR = Path("demo_output")
SAMPLE_PATH = DEMO_DIR / "report.docx"
RESULT_PATH = DEMO_DIR / "edited_report.docx"

PROMPT = (
    'insert a pie chart titled "Q2 Revenue" '
    "with Product 40, Services 35, Support 25 "
    "after the Financials heading"
)


def _build_sample_doc() -> None:
    """Create a minimal sample .docx for the demo."""
    doc = DocxDocument()
    doc.add_heading("Annual Report", level=1)
    doc.add_paragraph(
        "This report summarises company performance for the fiscal year."
    )
    doc.add_heading("Financials", level=1)
    doc.add_paragraph(
        "Q2 revenue exceeded expectations across all product lines."
    )
    doc.add_heading("Outlook", level=1)
    doc.add_paragraph("Growth is projected to continue into Q3 and Q4.")
    doc.save(str(SAMPLE_PATH))
    print(f"  Created sample document: {SAMPLE_PATH}")


def _print_outline(session: DocumentSession) -> None:
    outline = build_outline(session)
    print("\n  Document outline:")
    print(
        "\n".join(
            f"    {line}"
            for line in json.dumps(outline, indent=2).splitlines()
        )
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("=" * 60)
    print("  Document Editing Agent — offline demo")
    print("=" * 60)

    # Set up output directory
    DEMO_DIR.mkdir(exist_ok=True)

    # Build sample document
    print("\n[1/4] Building sample document …")
    _build_sample_doc()

    # Copy to a working path (mirrors what storage.py does for real jobs)
    working_path = DEMO_DIR / "working_report.docx"
    shutil.copy2(SAMPLE_PATH, working_path)

    # Open session and show outline
    print("\n[2/4] Analysing document structure …")
    session = DocumentSession(working_path)
    _print_outline(session)

    # Run the agent
    print(f"\n[3/4] Running agent …")
    print(f"  Prompt: {PROMPT!r}")
    print()

    changes = run_agent(
        session,
        PROMPT,
        on_change=lambda tool, desc: print(f"  ✓ [{tool}] {desc}"),
    )

    if not changes:
        print("  (no changes made)")
        sys.exit(1)

    # Save result
    print(f"\n[4/4] Saving result …")
    shutil.copy2(working_path, RESULT_PATH)
    working_path.unlink()  # clean up working copy

    print(f"  Output: {RESULT_PATH.resolve()}")
    print()
    print("  Change summary:")
    for i, change in enumerate(changes, 1):
        print(f"    {i}. [{change['tool']}] {change['description']}")

    print()
    print("=" * 60)
    print("  Demo complete. Open demo_output/edited_report.docx to verify.")
    print("=" * 60)


if __name__ == "__main__":
    main()
