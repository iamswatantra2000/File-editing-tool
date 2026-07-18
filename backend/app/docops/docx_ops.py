"""
python-docx operations.
All functions accept a DocumentSession and mutate the working copy in place,
then call session.save_docx() before returning.
"""

from pathlib import Path
from typing import List, Optional

from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.docops.session import DocumentSession


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _find_heading(session: DocumentSession, heading_text: str):
    """Return the paragraph whose text matches heading_text (case-insensitive)."""
    needle = heading_text.strip().lower()
    for para in session.docx.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip().lower() == needle:
            return para
    return None


def _find_paragraph(session: DocumentSession, text: str):
    """Return the first paragraph containing text (case-insensitive)."""
    needle = text.strip().lower()
    for para in session.docx.paragraphs:
        if needle in para.text.strip().lower():
            return para
    return None


def insert_paragraph_after(session: DocumentSession, anchor_text: str, new_text: str, style: str = "Normal") -> str:
    """Insert a new paragraph with new_text immediately after the anchor paragraph."""
    anchor = _find_paragraph(session, anchor_text)
    if anchor is None:
        raise ValueError(f"Anchor paragraph not found: {anchor_text!r}")

    new_para = session.docx.add_paragraph(new_text, style=style)
    # Move the new paragraph's XML node to immediately after the anchor
    anchor._p.addnext(new_para._p)
    session.save_docx()
    return f"Inserted paragraph after '{anchor_text}'"


def replace_text(session: DocumentSession, old_text: str, new_text: str) -> str:
    """Replace all occurrences of old_text with new_text across all paragraphs."""
    count = 0
    for para in session.docx.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    # Also check table cells
    for table in session.docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    session.save_docx()
    return f"Replaced {count} occurrence(s) of '{old_text}' with '{new_text}'"


def insert_image_after(
    session: DocumentSession,
    anchor_text: str,
    image_path: Path,
    width_inches: float = 4.5,
    caption: Optional[str] = None,
) -> str:
    """Insert an image (and optional caption) after the anchor paragraph."""
    anchor = _find_paragraph(session, anchor_text)
    if anchor is None:
        raise ValueError(f"Anchor paragraph not found: {anchor_text!r}")

    # Add image paragraph
    img_para = session.docx.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    anchor._p.addnext(img_para._p)

    if caption:
        cap_para = session.docx.add_paragraph(caption, style="Caption")
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para._p.addnext(cap_para._p)

    session.save_docx()
    return f"Inserted image after '{anchor_text}'" + (f" with caption '{caption}'" if caption else "")


def insert_table_after(
    session: DocumentSession,
    anchor_text: str,
    headers: List[str],
    rows: List[List[str]],
    style: str = "Table Grid",
) -> str:
    """Insert a table with headers and rows after the anchor paragraph."""
    anchor = _find_paragraph(session, anchor_text)
    if anchor is None:
        raise ValueError(f"Anchor paragraph not found: {anchor_text!r}")

    table = session.docx.add_table(rows=1 + len(rows), cols=len(headers), style=style)

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            table.cell(r_idx + 1, c_idx).text = str(value)

    anchor._p.addnext(table._tbl)
    session.save_docx()
    return f"Inserted {len(rows)}-row table after '{anchor_text}'"


def insert_heading_after(
    session: DocumentSession,
    anchor_text: str,
    heading_text: str,
    level: int = 1,
) -> str:
    """Insert a heading after the anchor paragraph."""
    anchor = _find_paragraph(session, anchor_text)
    if anchor is None:
        raise ValueError(f"Anchor paragraph not found: {anchor_text!r}")

    new_heading = session.docx.add_heading(heading_text, level=level)
    anchor._p.addnext(new_heading._p)
    session.save_docx()
    return f"Inserted Heading {level} '{heading_text}' after '{anchor_text}'"
