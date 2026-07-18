"""
Shared pytest fixtures.
"""

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook

from app.docops.session import DocumentSession


@pytest.fixture()
def docx_session(tmp_path: Path) -> DocumentSession:
    """A DocumentSession wrapping a minimal .docx with one heading and two paragraphs."""
    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_heading("Financials", level=1)
    doc.add_paragraph("Q2 revenue was strong.")

    path = tmp_path / "test.docx"
    doc.save(str(path))
    return DocumentSession(path)


@pytest.fixture()
def xlsx_session(tmp_path: Path) -> DocumentSession:
    """A DocumentSession wrapping a minimal .xlsx with one sheet and a header row."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Q1", "Q2", "Q3"])
    ws.append(["Product", 100, 200, 150])
    ws.append(["Services", 80, 130, 110])

    path = tmp_path / "test.xlsx"
    wb.save(str(path))
    return DocumentSession(path)
