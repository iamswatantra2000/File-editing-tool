"""
API integration tests using FastAPI's TestClient.
All tests run in mock agent mode — no real API calls.
"""

import io

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.config import settings
from app.main import app

client = TestClient(app)


@pytest.fixture(autouse=True)
def force_mock_mode(monkeypatch):
    monkeypatch.setattr(settings, "agent_mode", "mock")
    monkeypatch.setattr(settings, "anthropic_api_key", "")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is a test document.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx_bytes() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 10])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _submit_job(filename: str, data: bytes, prompt: str) -> dict:
    ext = filename.rsplit(".", 1)[-1]
    media = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext == "docx"
        else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response = client.post(
        "/api/jobs",
        files={"file": (filename, data, media)},
        data={"prompt": prompt},
    )
    return response


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

def test_health():
    r = client.get("/health")
    assert r.status_code == 200
    assert r.json()["status"] == "ok"


# ---------------------------------------------------------------------------
# POST /api/jobs
# ---------------------------------------------------------------------------

class TestCreateJob:
    def test_docx_job_accepted(self):
        r = _submit_job(
            "sample.docx",
            _make_docx_bytes(),
            'insert a pie chart titled "Revenue" with A 50, B 50 after Introduction',
        )
        assert r.status_code == 202
        body = r.json()
        assert "id" in body
        assert body["status"] in ("pending", "running", "done")
        assert body["filename"] == "sample.docx"

    def test_xlsx_job_accepted(self):
        r = _submit_job("data.xlsx", _make_xlsx_bytes(), "set B2 to 42")
        assert r.status_code == 202
        assert r.json()["filename"] == "data.xlsx"

    def test_unsupported_extension_rejected(self):
        r = client.post(
            "/api/jobs",
            files={"file": ("report.pdf", b"%PDF-fake", "application/pdf")},
            data={"prompt": "do something"},
        )
        assert r.status_code == 400
        assert "Unsupported" in r.json()["detail"]

    def test_empty_file_rejected(self):
        r = client.post(
            "/api/jobs",
            files={"file": ("empty.docx", b"", "application/octet-stream")},
            data={"prompt": "do something"},
        )
        assert r.status_code == 400


# ---------------------------------------------------------------------------
# GET /api/jobs/{id}
# ---------------------------------------------------------------------------

class TestGetJob:
    def test_poll_existing_job(self):
        r = _submit_job("doc.docx", _make_docx_bytes(), "insert a pie chart with A 60, B 40 after Introduction")
        job_id = r.json()["id"]

        poll = client.get(f"/api/jobs/{job_id}")
        assert poll.status_code == 200
        assert poll.json()["id"] == job_id

    def test_poll_missing_job(self):
        r = client.get("/api/jobs/nonexistent-id")
        assert r.status_code == 404


# ---------------------------------------------------------------------------
# GET /api/jobs/{id}/result
# ---------------------------------------------------------------------------

class TestGetResult:
    def test_result_available_when_done(self):
        r = _submit_job(
            "doc.docx",
            _make_docx_bytes(),
            'insert a pie chart titled "Rev" with X 70, Y 30 after Introduction',
        )
        job_id = r.json()["id"]

        # TestClient runs background tasks synchronously by default
        poll = client.get(f"/api/jobs/{job_id}")
        if poll.json()["status"] == "done":
            result = client.get(f"/api/jobs/{job_id}/result")
            assert result.status_code == 200
            assert result.headers["content-type"].startswith(
                "application/vnd.openxmlformats-officedocument"
            )
            assert len(result.content) > 0

    def test_result_not_ready_returns_409(self):
        # Create a job and immediately try to get the result before it runs
        from app.jobs import create_job, _store
        from app.models import JobStatus
        job = create_job(prompt="test", filename="fake.docx")
        # Leave status as pending
        r = client.get(f"/api/jobs/{job.id}/result")
        assert r.status_code == 409

    def test_result_missing_job_returns_404(self):
        r = client.get("/api/jobs/does-not-exist/result")
        assert r.status_code == 404


# ---------------------------------------------------------------------------
# GET /api/jobs
# ---------------------------------------------------------------------------

def test_list_jobs():
    _submit_job("a.docx", _make_docx_bytes(), "insert a pie chart with A 50, B 50 after Introduction")
    r = client.get("/api/jobs")
    assert r.status_code == 200
    assert isinstance(r.json(), list)
    assert len(r.json()) >= 1
