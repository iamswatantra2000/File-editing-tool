"""
Unit tests for docops operations.
"""

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.docops import docx_ops, xlsx_ops, charts
from app.docops.outline import build_outline
from app.docops.session import DocumentSession


# ---------------------------------------------------------------------------
# outline
# ---------------------------------------------------------------------------

class TestOutline:
    def test_docx_outline_headings(self, docx_session):
        outline = build_outline(docx_session)
        assert outline["type"] == "docx"
        headings = [h["text"] for h in outline["headings"]]
        assert "Introduction" in headings
        assert "Financials" in headings

    def test_docx_outline_paragraphs(self, docx_session):
        outline = build_outline(docx_session)
        assert any("first paragraph" in p for p in outline["paragraphs"])

    def test_xlsx_outline_sheets(self, xlsx_session):
        outline = build_outline(xlsx_session)
        assert outline["type"] == "xlsx"
        assert outline["sheets"][0]["name"] == "Sheet1"
        assert outline["sheets"][0]["max_row"] >= 3


# ---------------------------------------------------------------------------
# docx_ops
# ---------------------------------------------------------------------------

class TestDocxOps:
    def test_replace_text(self, docx_session):
        result = docx_ops.replace_text(docx_session, "first paragraph", "opening paragraph")
        assert "1" in result or "occurrence" in result
        # Verify change persisted on disk
        reloaded = DocxDocument(str(docx_session.path))
        texts = [p.text for p in reloaded.paragraphs]
        assert any("opening paragraph" in t for t in texts)
        assert not any("first paragraph" in t for t in texts)

    def test_replace_text_no_match(self, docx_session):
        result = docx_ops.replace_text(docx_session, "nonexistent string", "replacement")
        assert "0" in result

    def test_insert_paragraph_after(self, docx_session):
        docx_ops.insert_paragraph_after(
            docx_session,
            anchor_text="first paragraph",
            new_text="Inserted paragraph.",
        )
        reloaded = DocxDocument(str(docx_session.path))
        texts = [p.text for p in reloaded.paragraphs]
        assert "Inserted paragraph." in texts

    def test_insert_paragraph_bad_anchor(self, docx_session):
        with pytest.raises(ValueError, match="not found"):
            docx_ops.insert_paragraph_after(docx_session, "DOES NOT EXIST", "text")

    def test_insert_heading_after(self, docx_session):
        docx_ops.insert_heading_after(
            docx_session,
            anchor_text="Introduction",
            heading_text="New Section",
            level=2,
        )
        reloaded = DocxDocument(str(docx_session.path))
        headings = [p.text for p in reloaded.paragraphs if p.style.name.startswith("Heading")]
        assert "New Section" in headings

    def test_insert_table_after(self, docx_session):
        docx_ops.insert_table_after(
            docx_session,
            anchor_text="Q2 revenue",
            headers=["Item", "Value"],
            rows=[["Product", "40"], ["Services", "35"]],
        )
        reloaded = DocxDocument(str(docx_session.path))
        assert len(reloaded.tables) == 1
        assert reloaded.tables[0].cell(0, 0).text == "Item"

    def test_insert_image_after(self, docx_session, tmp_path):
        # Generate a real chart PNG first
        chart_path = charts.pie_chart(
            labels=["A", "B", "C"],
            values=[30, 40, 30],
            title="Test Chart",
        )
        docx_ops.insert_image_after(
            docx_session,
            anchor_text="Introduction",
            image_path=chart_path,
            caption="Figure 1",
        )
        # Validate the file can be reopened
        reloaded = DocxDocument(str(docx_session.path))
        assert reloaded is not None


# ---------------------------------------------------------------------------
# xlsx_ops
# ---------------------------------------------------------------------------

class TestXlsxOps:
    def test_write_cell(self, xlsx_session):
        result = xlsx_ops.write_cell(xlsx_session, "Sheet1", "B2", 999)
        assert "999" in result
        xlsx_session.reload()
        assert xlsx_session.workbook["Sheet1"]["B2"].value == 999

    def test_read_cell(self, xlsx_session):
        val = xlsx_ops.read_cell(xlsx_session, "Sheet1", "A1")
        assert val == "Name"

    def test_append_row(self, xlsx_session):
        xlsx_ops.append_row(xlsx_session, "Sheet1", ["Support", 50, 60, 70])
        xlsx_session.reload()
        ws = xlsx_session.workbook["Sheet1"]
        last_row = [ws.cell(row=ws.max_row, column=c).value for c in range(1, 5)]
        assert last_row[0] == "Support"

    def test_write_range(self, xlsx_session):
        xlsx_ops.write_range(
            xlsx_session,
            "Sheet1",
            "A5",
            [["X", 1, 2, 3], ["Y", 4, 5, 6]],
        )
        xlsx_session.reload()
        ws = xlsx_session.workbook["Sheet1"]
        assert ws["A5"].value == "X"
        assert ws["A6"].value == "Y"

    def test_edit_table_cell(self, xlsx_session):
        xlsx_ops.edit_table_cell(xlsx_session, "Sheet1", row=2, col=2, value=42)
        xlsx_session.reload()
        assert xlsx_session.workbook["Sheet1"].cell(row=2, column=2).value == 42

    def test_add_sheet(self, xlsx_session):
        xlsx_ops.add_sheet(xlsx_session, "NewSheet")
        xlsx_session.reload()
        assert "NewSheet" in xlsx_session.workbook.sheetnames

    def test_add_sheet_duplicate_raises(self, xlsx_session):
        with pytest.raises(ValueError, match="already exists"):
            xlsx_ops.add_sheet(xlsx_session, "Sheet1")

    def test_rename_sheet(self, xlsx_session):
        xlsx_ops.rename_sheet(xlsx_session, "Sheet1", "Renamed")
        xlsx_session.reload()
        assert "Renamed" in xlsx_session.workbook.sheetnames
        assert "Sheet1" not in xlsx_session.workbook.sheetnames

    def test_delete_sheet(self, xlsx_session):
        xlsx_ops.add_sheet(xlsx_session, "Temp")
        xlsx_ops.delete_sheet(xlsx_session, "Temp")
        xlsx_session.reload()
        assert "Temp" not in xlsx_session.workbook.sheetnames


# ---------------------------------------------------------------------------
# charts
# ---------------------------------------------------------------------------

class TestCharts:
    def test_pie_chart_creates_png(self):
        path = charts.pie_chart(
            labels=["A", "B", "C"],
            values=[10, 20, 70],
            title="Revenue",
        )
        assert path.exists()
        assert path.suffix == ".png"
        assert path.stat().st_size > 0

    def test_bar_chart_creates_png(self):
        path = charts.bar_chart(
            labels=["Jan", "Feb", "Mar"],
            values=[100, 150, 120],
            title="Monthly",
            ylabel="Sales",
        )
        assert path.exists()
        assert path.suffix == ".png"

    def test_build_chart_unknown_type(self):
        with pytest.raises(ValueError, match="Unknown chart type"):
            charts.build_chart("scatter", labels=[], values=[])
